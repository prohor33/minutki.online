"""
библиотеки: pip install streamlit numpy docx python-docx
запуск приложения: STREAMLIT_SERVER_PORT=80 streamlit run app.py --server.maxUploadSize=1028
"""

import streamlit as st
import pandas as pd
import subprocess as sp
import shlex
import json
import os
import base64
from docx import Document
 
BUTTON_STYLE = (
    "-webkit-border-radius: 4px; "
    "-moz-border-radius: 4px; "
    "border-radius: 4px; "
    "border: solid 1px #20538D; "
    "text-shadow: 0 -1px 0 rgba(0, 0, 0, 0.4); "
    "-webkit-box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.4), 0 1px 1px rgba(0, 0, 0, 0.2); "
    "-moz-box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.4), 0 1px 1px rgba(0, 0, 0, 0.2); "
    "box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.4), 0 1px 1px rgba(0, 0, 0, 0.2); "
    "background: #4479BA; "
    "color: #FFF; "
    "padding: 8px 12px; "
    "text-decoration: none;"
)


def generate_report(json_obj, output_path):
    # with open(input_path) as fp:
    #     output = json.load(fp)
    output = json_obj
    document = Document()
    document.add_heading('Полная транскрибация записи совещания.')
    document.add_paragraph(output["text"])
    document.add_heading('Поручения по итогу совещания.')
    # таблица с карточками
    rows = len(output["cards"])
    cols = len(output["cards"][0]) or 0
    table = document.add_table(rows=rows, cols=cols)
    for icard, row in enumerate(table.rows):
        for col_name, cell in zip(output["cards"][icard], row.cells):
            cell.text = output["cards"][icard][col_name]["text"]
    document.save(output_path)

st.set_page_config(layout="wide")
uploaded_files = st.file_uploader(label='Выберите файл для загрузки', type="MP4")
 
if uploaded_files is not None:
    video_bytes = uploaded_files.read()
    st.video(video_bytes)
    with st.spinner('Пожалуйста подождите, идет анализ данных...'):
        normal = sp.run(
            shlex.split("python runner.py"),
            stdout=sp.PIPE, stderr=sp.PIPE,
            check=True,
            text=True
        )
        output = normal.stdout
        obj = json.loads(output)
        
        # разбиваем на две колонки
        with st.container():
            left, right = st.columns(2)
            with left:
                st.markdown("<h2 style='text-align: center;'>Текст</h2>", unsafe_allow_html=True)
                st.write(obj["text"])
            with right:
                st.markdown("<h2 style='text-align: center;'>Поручения</h2>", unsafe_allow_html=True)
                st.markdown("---")
            for record in obj["cards"]:
                with right:
                    st.write(record["assignment"]["text"])
                    st.caption(f"Ответственный: {record['responsible']['text']}")
                    st.caption(f"Срок: {record['date']['text']}")
                    st.markdown("---")
        
        # генерируем отчет
        report_path = "report/Отчёт.docx"
        generate_report(json_obj=obj, output_path=report_path)
        with open(report_path, "rb") as fp:
            bytes = fp.read()
            b64 = base64.b64encode(bytes).decode()
            href = f'<a align="center" style="{BUTTON_STYLE}" href="data:base64,{b64}" download="{os.path.basename(report_path)}">Скачать отчет в формате Word</a>'
            st.markdown(href, unsafe_allow_html=True)
