from docx import Document
import json

output = json.load(open("test.json"))

document = Document()
document.add_heading('Полная транскрибация записи совещания.')
document.add_paragraph(output["text"])
document.add_heading('Поручения по итогу совещания.')
# таблица с карточками
rows = len(output["cards"])
cols = len(output["cards"][0]) or 0
table = document.add_table(rows=rows, cols=cols)
for icard, row in enumerate(table.rows):
    for col_name, cell in zip(output["cards"][icard], row.cells):
        cell.text = output["cards"][icard][col_name]["text"]
document.save('test_.docx')

